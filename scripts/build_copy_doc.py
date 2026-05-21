#!/usr/bin/env python3
"""Build a labeled Word doc of all site copy for a non-technical editor.

Pipeline:
  1. Runs scripts/export-copy.ts (Node type-stripping) to dump the real
     source data (portfolio.ts + site.ts) to JSON — so the doc stays accurate
     as the site changes.
  2. Renders that JSON into a structured .docx with python-docx.
  3. Saves to ~/Downloads with a timestamp.

Usage:  python3 scripts/build_copy_doc.py
Requires: node >=22 (for --experimental-strip-types), python-docx.

NOTE on "shown vs. not shown": some copy lives in the data but is currently
commented out in the components (hero meta block, the Outcome phase, the
Next-project link). Those are labeled "NOT CURRENTLY SHOWN ON SITE" so the
editor has full context without guessing.
"""

import json
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
DOWNLOADS = Path.home() / "Downloads"

GREY = RGBColor(0x80, 0x80, 0x80)
ACCENT = RGBColor(0x00, 0x66, 0xCC)
MUTED_NOTE = RGBColor(0xB0, 0x6A, 0x00)


def strip_html(s: str) -> str:
    """Flatten the small amount of inline HTML used in copy fields.
    <br/> -> newline, <em>/<strong>/<p> stripped, entities decoded."""
    if not s:
        return ""
    s = re.sub(r"<\s*br\s*/?\s*>", "\n", s, flags=re.I)
    s = re.sub(r"</\s*p\s*>", "\n", s, flags=re.I)
    s = re.sub(r"<[^>]+>", "", s)
    s = s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    return s.strip()


def export_json() -> dict:
    tmp = Path(tempfile.gettempdir()) / "site-copy-export.json"
    subprocess.run(
        [
            "node",
            "--experimental-strip-types",
            str(REPO / "scripts" / "export-copy.ts"),
            str(tmp),
        ],
        check=True,
        cwd=str(REPO),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return json.loads(tmp.read_text())


# ---------- doc helpers ----------

def label(doc, text):
    """A small grey 'what this is' caption above a piece of copy."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    r.font.all_caps = True
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    return p


def copy_block(doc, text, italic=False):
    """The actual editable copy."""
    text = text or ""
    for i, line in enumerate(text.split("\n")):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.italic = italic
        p.paragraph_format.space_after = Pt(2)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("⚠ " + text)
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED_NOTE
    r.italic = True
    return p


def field(doc, name, value, italic=False):
    label(doc, name)
    copy_block(doc, value, italic=italic)


def bullets(doc, items):
    for pair in items:
        lab, txt = pair[0], pair[1]
        p = doc.add_paragraph(style="List Bullet")
        rl = p.add_run(f"{lab}: ")
        rl.bold = True
        p.add_run(strip_html(txt))


# ---------- build ----------

def build(data: dict) -> Path:
    doc = Document()

    # Title
    h = doc.add_heading("adrianbarn.es — Site Copy", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run(
        "All editable copy for the portfolio site, labeled by location. "
        "Generated " + datetime.now().strftime("%Y-%m-%d %H:%M")
    )
    r.font.color.rgb = GREY
    r.italic = True
    doc.add_paragraph(
        "Edit the text under each grey label. Items marked NOT CURRENTLY "
        "SHOWN exist in the data but are hidden on the live site right now."
    )

    # --- SITE / SEO ---
    doc.add_heading("Browser / SEO Metadata", level=1)
    note(doc, "Shown in the browser tab and search-engine results, not on the page itself.")
    site = data["SITE"]
    field(doc, "Site name", site.get("name", ""))
    field(doc, "Page title (browser tab + Google result title)", site.get("title", ""))
    field(doc, "Meta description (Google result snippet)", site.get("description", ""))

    # --- LANDER / INTRO ---
    doc.add_heading("Landing / Intro Section (top of page)", level=1)
    lander = data["LANDER"]
    field(doc, "Wordmark (top-left logo text)", lander.get("mark", ""))
    field(doc, "Kicker (small label above name)", lander.get("kicker", ""))
    field(doc, "Name — line 1", lander.get("nameTop", ""))
    field(doc, "Name — line 2", lander.get("nameBottom", ""))
    field(doc, "Intro blurb (main paragraph)", strip_html(lander.get("blurb", "")))
    doc.add_heading("Intro — fact row (bottom of intro)", level=2)
    for item in lander.get("bottom", []):
        field(doc, item.get("label", ""), item.get("value", ""))
    doc.add_heading("Intro — navigation links", level=2)
    for item in lander.get("nav", []):
        suffix = "  (external link)" if item.get("external") else ""
        field(doc, "Nav link label" + suffix, item.get("label", ""))

    # --- CASE STUDIES ---
    doc.add_heading("Case Studies", level=1)
    note(doc, "Each case study has a HERO (full-screen intro) and a BODY (the detail sections below it).")

    PHASE_LABELS = [
        ("problem", "Phase 1"),
        ("approach", "Phase 2"),
        ("built", "Phase 3"),
    ]

    for cs in data["CASE_STUDIES"]:
        doc.add_page_break()
        doc.add_heading(f"{cs['num']} · {cs['company']}", level=1)

        # Hero
        doc.add_heading("Hero (full-screen intro)", level=2)
        field(doc, "Company name (kicker)", cs.get("company", ""))
        field(doc, "Short name (used in side nav)", cs.get("shortName", ""))
        field(doc, "Hero title (large headline — italics shown in italic)",
              strip_html(cs.get("titleHtml", "")),
              italic=("<em>" in cs.get("titleHtml", "")))
        field(doc, "Hero summary (sentence under the title)", cs.get("summary", ""))

        # Hero meta — currently hidden
        doc.add_heading("Hero meta facts", level=3)
        note(doc, "NOT CURRENTLY SHOWN ON SITE (hero meta block is commented out), but kept in data.")
        for lab, key in [("Role", "role"), ("Duration", "duration"),
                         ("Team", "team"), ("Years", "year"),
                         ("Location", "location")]:
            field(doc, lab, cs.get(key, ""))

        # Body phases
        doc.add_heading("Body — detail sections", level=2)
        for key, fallback_label in PHASE_LABELS:
            phase = cs.get(key)
            if not phase:
                continue
            doc.add_heading(
                f"{fallback_label}  ({key})", level=3
            )
            field(doc, "Section number/label (small, above heading)", phase.get("num", ""))
            field(doc, "Section heading", phase.get("name", ""))
            field(doc, "Section paragraph", strip_html(phase.get("prose", "")))
            if phase.get("bullets"):
                label(doc, "Section bullets (label: text)")
                bullets(doc, phase["bullets"])

        # Outcome — currently hidden
        outcome = cs.get("outcome")
        if outcome:
            doc.add_heading("Outcome + metrics", level=3)
            note(doc, "NOT CURRENTLY SHOWN ON SITE (Outcome phase is commented out), but kept in data.")
            field(doc, "Outcome paragraph", strip_html(outcome.get("prose", "")))
            label(doc, "Metrics (value — label)")
            for m in outcome.get("metrics", []):
                p = doc.add_paragraph(style="List Bullet")
                rb = p.add_run(f"{m.get('value','')} — ")
                rb.bold = True
                p.add_run(m.get("label", ""))

        # Stack
        if cs.get("stack"):
            label(doc, "Stack & Tools (pills shown at the bottom of the body)")
            doc.add_paragraph(" · ".join(cs["stack"]))

        # Gallery captions
        gallery = cs.get("gallery", [])
        caand = [g for g in gallery if g.get("caption")]
        if caand:
            doc.add_heading("Carousel / gallery captions", level=3)
            note(doc, "One caption per media slide (image/video). Order matches the carousel.")
            for g in gallery:
                cap = g.get("caption")
                if not cap:
                    continue
                pos = g.get("gallery", "")
                field(doc, f"{g.get('type','media')} caption ({pos} carousel)", cap)

    # --- CONTACT ---
    doc.add_page_break()
    doc.add_heading("Contact Section (bottom of page)", level=1)
    contact = data["CONTACT"]
    field(doc, "Kicker (small label)", strip_html(contact.get("kicker", "")))
    field(doc, "Headline", strip_html(contact.get("headline", "")))
    field(doc, "Email", contact.get("email", ""))
    field(doc, "Phone (display)", contact.get("phoneDisplay", ""))
    field(doc, "LinkedIn handle", contact.get("linkedinHandle", ""))
    field(doc, "Location", contact.get("location", ""))
    if contact.get("colophon"):
        field(doc, "Colophon (fine print)", contact.get("colophon", ""))

    # --- Fixed UI labels ---
    doc.add_heading("Fixed UI Labels", level=1)
    note(doc, "Structural words baked into the layout. Editable but rarely changed.")
    for lab in ["Case Study (eyebrow above each hero number)",
                "Stack & Tools (label above the tech pills)"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(lab)

    ts = datetime.now().strftime("%Y-%m-%d_%H%M")
    out = DOWNLOADS / f"adrianbarnes-site-copy_{ts}.docx"
    doc.save(str(out))
    return out


if __name__ == "__main__":
    data = export_json()
    out = build(data)
    print(str(out))
